"""Markdown → docx/pdf 导出（防回归：docx 主循环须递增行指针）。"""
from __future__ import annotations

import io

from django.test import SimpleTestCase
from docx import Document

from pipeline.reporting.md_document_export import (
    markdown_to_docx_bytes,
    markdown_to_pdf_bytes,
)


class MdDocumentExportTests(SimpleTestCase):
    def test_docx_plain_lines_terminate(self) -> None:
        md = "第一行\n\n第二行\n仍是一段"
        data = markdown_to_docx_bytes(md)
        self.assertGreater(len(data), 2000)
        self.assertTrue(data.startswith(b"PK"))

    def test_pdf_plain_lines_terminate(self) -> None:
        md = "标题\n\n正文一行"
        data = markdown_to_pdf_bytes(md)
        self.assertGreater(len(data), 100)
        self.assertTrue(data.startswith(b"%PDF"))

    def test_docx_merges_soft_line_breaks_in_paragraph(self) -> None:
        """模型/编辑器折行不应被当成多个独立段落。"""
        md = "统一商详第\n1 屏核心话术\n\n下一段"
        data = markdown_to_docx_bytes(md)
        doc = Document(io.BytesIO(data))
        texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        self.assertEqual(texts, ["统一商详第 1 屏核心话术", "下一段"])

    def test_docx_task_list_strips_prefix_uses_normal_bullet(self) -> None:
        """``- [x]`` / ``- [ ]`` 去掉方括号标记，用普通列表符号，观感接近 MD 预览。"""
        md = "- [x] 卡腰：围绕中位。\n- [ ] 贴顶：高位。\n- 普通列表项"
        data = markdown_to_docx_bytes(md)
        doc = Document(io.BytesIO(data))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        joined = "\n".join(texts)
        self.assertNotIn("☑", joined)
        self.assertNotIn("☐", joined)
        self.assertNotIn("[x]", joined)
        self.assertNotIn("[ ]", joined)
        self.assertIn("卡腰：围绕中位", joined)
        self.assertIn("贴顶：高位", joined)
        self.assertTrue(any("普通列表项" in t for t in texts))

    def test_docx_table_skips_long_dash_separator_row(self) -> None:
        """`|----------|` 分隔行不得作为表格正文行导出。"""
        md = "| 差异点 | 说明 |\n|----------|----------|\n| A | B |"
        data = markdown_to_docx_bytes(md)
        doc = Document(io.BytesIO(data))
        self.assertEqual(len(doc.tables), 1)
        self.assertEqual(len(doc.tables[0].rows), 2)
        self.assertNotIn("----------", doc.tables[0].rows[1].cells[0].text)
