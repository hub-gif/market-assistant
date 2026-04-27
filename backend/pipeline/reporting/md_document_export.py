"""Markdown → Word（.docx）/ 简易 PDF；供任务报告与策略稿导出。"""
from __future__ import annotations

import os
import re
from io import BytesIO
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape as xml_escape


def _strip_inline_md(s: str) -> str:
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"`([^`]+)`", r"\1", s)
    return s


_RE_TASK_CHECKED = re.compile(r"^\[x\]\s*", re.IGNORECASE)
_RE_TASK_UNCHECKED = re.compile(r"^\[ \]\s*")
# GFM 表头分隔行：每格为 :--- / ---: / :---: / ---------- 等（至少 3 个连字符）
_TABLE_SEP_CELL = re.compile(r"^:?-{3,}:?$")


def _strip_gfm_task_list_prefix(text: str) -> str:
    """去掉 ``- [x]`` / ``- [ ]`` 中的任务标记，导出时用普通项目符号，观感接近 MD 预览圆点。"""
    t = text.strip()
    m = _RE_TASK_CHECKED.match(t)
    if m:
        return t[m.end() :].strip()
    m = _RE_TASK_UNCHECKED.match(t)
    if m:
        return t[m.end() :].strip()
    return text


def _is_table_sep(line: str) -> bool:
    """GFM 表头与表体之间的分隔行（含任意长度连字符，如 ``|----------|``）。"""
    row_line = line.strip()
    if not row_line.startswith("|"):
        return False
    cells = [c.strip() for c in row_line.strip("|").split("|")]
    sep_cells = [c for c in cells if c]
    if len(sep_cells) < 2:
        return False
    return all(_TABLE_SEP_CELL.match(c) is not None for c in sep_cells)


_RE_HEADING = re.compile(r"^(#{1,6})\s+(.+)$")
_RE_UL = re.compile(r"^\s*[-*+]\s+(.+)$")
_RE_OL = re.compile(r"^\s*(\d+)\.\s+(.+)$")
_RE_BLOCKQUOTE = re.compile(r"^\s*>\s?(.*)$")
_RE_HR = re.compile(r"^\s*(?:[-*_]\s*){3,}\s*$")

_img_line = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")


def _join_md_soft_break_lines(lines: list[str]) -> str:
    """把编辑器/模型折行产生的多行合并为一段（等价于 CommonMark 软换行 → 空格）。"""
    parts = [ln.strip() for ln in lines if ln and ln.strip()]
    if not parts:
        return ""
    return " ".join(parts)


def _is_plain_markdown_line(s: str) -> bool:
    """是否可作为「正文折行」参与合并的一行（非标题/列表/表格等）。"""
    t = s.strip()
    if not t:
        return False
    if t.startswith("```"):
        return False
    if _RE_HR.match(s):
        return False
    if _match_heading(s) is not None:
        return False
    if _img_line.match(t):
        return False
    if t.startswith("|"):
        return False
    if _RE_UL.match(s):
        return False
    if _RE_OL.match(s):
        return False
    if _RE_BLOCKQUOTE.match(s):
        return False
    return True


def _match_heading(line: str) -> tuple[int, str] | None:
    """返回 (docx level 0–8, 标题文本) 或 None。"""
    m = _RE_HEADING.match(line.strip())
    if not m:
        return None
    depth = len(m.group(1))
    title = _strip_inline_md(m.group(2).strip())
    level = min(max(depth - 1, 0), 8)
    return (level, title)


def markdown_to_docx_bytes(md: str, *, asset_root: Path | None = None) -> bytes:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Inches, Pt

    doc = Document()
    try:
        style = doc.styles["Normal"]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(10.5)
    except Exception:
        pass

    def _add_list_bullet(text: str) -> None:
        t = _strip_gfm_task_list_prefix(_strip_inline_md(text))
        try:
            doc.add_paragraph(t, style="List Bullet")
        except KeyError:
            doc.add_paragraph("• " + t)

    def _add_list_number(text: str) -> None:
        t = _strip_inline_md(text)
        try:
            doc.add_paragraph(t, style="List Number")
        except KeyError:
            doc.add_paragraph(t)

    lines = (md or "").replace("\r\n", "\n").split("\n")
    i = 0
    in_fence = False
    plain_buf: list[str] = []

    def flush_plain() -> None:
        if not plain_buf:
            return
        merged = _join_md_soft_break_lines(plain_buf)
        plain_buf.clear()
        if not merged:
            return
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.add_run(_strip_inline_md(merged))

    while i < len(lines):
        raw = lines[i]
        if raw.strip().startswith("```"):
            if not in_fence:
                flush_plain()
            in_fence = not in_fence
            i += 1
            continue
        if in_fence:
            p = doc.add_paragraph(xml_escape(raw) or " ")
            p.style = doc.styles["Normal"]
            for run in p.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            i += 1
            continue

        line = raw.rstrip()
        if not line.strip():
            flush_plain()
            doc.add_paragraph("")
            i += 1
            continue

        if _RE_HR.match(line):
            flush_plain()
            doc.add_paragraph("")
            i += 1
            continue

        hm = _match_heading(line)
        if hm is not None:
            flush_plain()
            doc.add_heading(hm[1], level=hm[0])
            i += 1
            continue

        mimg = _img_line.match(line.strip())
        if mimg and asset_root is not None:
            flush_plain()
            rel = mimg.group(2).strip()
            if not (rel.startswith("http://") or rel.startswith("https://")):
                img_path = (asset_root / rel).resolve()
                try:
                    img_path.relative_to(asset_root.resolve())
                except ValueError:
                    i += 1
                    continue
                if img_path.is_file():
                    doc.add_picture(str(img_path), width=Inches(5.9))
            i += 1
            continue

        if line.strip().startswith("|"):
            flush_plain()
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_line = lines[i].strip()
                if _is_table_sep(row_line):
                    i += 1
                    continue
                cells = [c.strip() for c in row_line.strip("|").split("|")]
                rows.append([_strip_inline_md(c) for c in cells])
                i += 1
            if rows:
                max_cols = max(len(r) for r in rows)
                pad_rows = [r + [""] * (max_cols - len(r)) for r in rows]
                tbl = doc.add_table(rows=len(pad_rows), cols=max_cols)
                tbl.style = "Table Grid"
                for ri, row in enumerate(pad_rows):
                    for ci, cell in enumerate(row):
                        tbl.rows[ri].cells[ci].text = cell
            continue

        mu = _RE_UL.match(line)
        if mu:
            flush_plain()
            _add_list_bullet(mu.group(1))
            i += 1
            continue

        mo = _RE_OL.match(line)
        if mo:
            flush_plain()
            _add_list_number(mo.group(2))
            i += 1
            continue

        mq = _RE_BLOCKQUOTE.match(line)
        if mq:
            flush_plain()
            inner = mq.group(1).strip()
            if inner:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.add_run(_strip_inline_md(inner))
            i += 1
            continue

        if _is_plain_markdown_line(line):
            plain_buf.append(line)
            i += 1
            continue

        flush_plain()
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.add_run(_strip_inline_md(line))
        i += 1

    flush_plain()
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _pdf_font_candidates() -> list[Path]:
    raw = (os.environ.get("MA_PDF_FONT") or "").strip()
    out: list[Path] = []
    if raw:
        out.append(Path(raw))
    windir = os.environ.get("WINDIR", r"C:\Windows")
    out.extend(
        [
            Path(windir) / "Fonts" / "simhei.ttf",
            Path(windir) / "Fonts" / "simsun.ttc",
            Path(windir) / "Fonts" / "msyh.ttf",
        ]
    )
    # Linux / 容器常见中文字体（路径不存在则跳过）
    out.extend(
        [
            Path("/usr/share/fonts/truetype/wqy/wqy-microhei.ttc"),
            Path("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"),
            Path("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc"),
            Path("/usr/share/fonts/truetype/noto/NotoSansCJKsc-Regular.otf"),
            Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
        ]
    )
    return out


def _pdf_flowable_image(img_path: Path, *, max_w: float, max_h: float) -> Any:
    """将插图缩放到不超过 max_w×max_h（ReportLab 单位，与 cm 一致），保持宽高比，避免矩阵长图撑爆版面。"""
    from reportlab.lib.utils import ImageReader
    from reportlab.platypus import Image as RLImage

    p = str(img_path)
    try:
        ir = ImageReader(p)
        iw, ih = ir.getSize()
    except Exception:
        return RLImage(p, width=max_w * 0.9, height=max_h * 0.9)
    if iw <= 0 or ih <= 0:
        return RLImage(p, width=max_w * 0.9, height=max_h * 0.9)
    w = float(max_w)
    h = w * (float(ih) / float(iw))
    if h > float(max_h):
        h = float(max_h)
        w = h * (float(iw) / float(ih))
    return RLImage(p, width=w, height=h)


def markdown_to_pdf_bytes(md: str, *, asset_root: Path | None = None) -> bytes:
    """简易 PDF；需本机 .ttf 中文字体或环境变量 MA_PDF_FONT。"""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    font_name = "MaExportCJK"
    registered = False
    for p in _pdf_font_candidates():
        if not p.is_file():
            continue
        try:
            if p.suffix.lower() == ".ttc":
                try:
                    pdfmetrics.registerFont(
                        TTFont(font_name, str(p), subfontIndex=0)
                    )
                except TypeError:
                    pdfmetrics.registerFont(TTFont(font_name, str(p)))
            else:
                pdfmetrics.registerFont(TTFont(font_name, str(p)))
            registered = True
            break
        except Exception:
            continue
    if not registered:
        raise ValueError(
            "未找到可用的中文字体文件。请在 Windows 上安装黑体/宋体，"
            "或设置环境变量 MA_PDF_FONT 指向 .ttf 文件路径。"
        )

    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        name="BodyCJK",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=10,
        leading=14,
        spaceAfter=3,
    )
    h1s = ParagraphStyle(
        name="H1CJK",
        parent=body,
        fontSize=16,
        leading=20,
        spaceBefore=0,
        spaceAfter=10,
    )
    h2s = ParagraphStyle(
        name="H2CJK",
        parent=body,
        fontSize=13,
        leading=17,
        spaceBefore=14,
        spaceAfter=6,
    )
    h3s = ParagraphStyle(
        name="H3CJK",
        parent=body,
        fontSize=12,
        leading=16,
        spaceBefore=10,
        spaceAfter=5,
    )
    h4s = ParagraphStyle(
        name="H4CJK",
        parent=body,
        fontSize=11,
        leading=15,
        spaceBefore=8,
        spaceAfter=4,
    )
    h56s = ParagraphStyle(
        name="H56CJK",
        parent=body,
        fontSize=10.5,
        leading=14,
        spaceBefore=6,
        spaceAfter=3,
    )
    quote_style = ParagraphStyle(
        name="QuoteCJK",
        parent=body,
        leftIndent=14,
        fontSize=9.5,
        textColor=colors.HexColor("#444444"),
    )
    # 项目符号用 Helvetica 绘制：正文 CJK 字体常缺 U+2022「•」，会落成方框（似 ☐）
    bullet_body = ParagraphStyle(
        name="BulletBodyCJK",
        parent=body,
        leftIndent=22,
        bulletIndent=10,
        firstLineIndent=0,
        bulletFontName="Helvetica",
        bulletFontSize=10,
        wordWrap="CJK",
    )
    ol_body = ParagraphStyle(
        name="OlBodyCJK",
        parent=body,
        leftIndent=18,
        firstLineIndent=0,
        wordWrap="CJK",
    )

    story: list[Any] = []
    lines = (md or "").replace("\r\n", "\n").split("\n")
    i = 0
    in_fence = False
    plain_buf: list[str] = []

    def _para_cell(s: str, style: Any) -> Paragraph:
        return Paragraph(xml_escape(_strip_inline_md(s)), style)

    def flush_plain_pdf() -> None:
        if not plain_buf:
            return
        merged = _join_md_soft_break_lines(plain_buf)
        plain_buf.clear()
        if not merged:
            return
        story.append(
            Paragraph(xml_escape(_strip_inline_md(merged)), body)
        )

    while i < len(lines):
        raw = lines[i]
        if raw.strip().startswith("```"):
            if not in_fence:
                flush_plain_pdf()
            in_fence = not in_fence
            i += 1
            continue
        s = raw.rstrip()
        if in_fence:
            story.append(Paragraph(xml_escape(s or " "), body))
            story.append(Spacer(1, 0.1 * cm))
            i += 1
            continue
        if not s.strip():
            flush_plain_pdf()
            story.append(Spacer(1, 0.15 * cm))
            i += 1
            continue

        if _RE_HR.match(s):
            flush_plain_pdf()
            story.append(Spacer(1, 0.2 * cm))
            i += 1
            continue

        hm = _match_heading(s)
        if hm is not None:
            flush_plain_pdf()
            level, title = hm
            title_esc = xml_escape(title)
            if level == 0:
                story.append(Paragraph(title_esc, h1s))
            elif level == 1:
                story.append(Paragraph(title_esc, h2s))
            elif level == 2:
                story.append(Paragraph(title_esc, h3s))
            elif level == 3:
                story.append(Paragraph(title_esc, h4s))
            else:
                story.append(Paragraph(title_esc, h56s))
            i += 1
            continue

        mimg = _img_line.match(s.strip())
        if mimg and asset_root is not None:
            flush_plain_pdf()
            rel = mimg.group(2).strip()
            if not (rel.startswith("http://") or rel.startswith("https://")):
                img_path = (asset_root / rel).resolve()
                try:
                    img_path.relative_to(asset_root.resolve())
                except ValueError:
                    i += 1
                    continue
                if img_path.is_file():
                    story.append(
                        _pdf_flowable_image(
                            img_path, max_w=13 * cm, max_h=24 * cm
                        )
                    )
                    story.append(Spacer(1, 0.2 * cm))
            i += 1
            continue

        if s.strip().startswith("|"):
            flush_plain_pdf()
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_line = lines[i].strip()
                if _is_table_sep(row_line):
                    i += 1
                    continue
                cells = [c.strip() for c in row_line.strip("|").split("|")]
                rows.append([_strip_inline_md(c) for c in cells])
                i += 1
            if rows:
                max_cols = max(len(r) for r in rows)
                pad_rows = [r + [""] * (max_cols - len(r)) for r in rows]
                usable_w = 17 * cm
                if max_cols == 2:
                    col_widths = [4.2 * cm, usable_w - 4.2 * cm]
                else:
                    col_widths = [usable_w / float(max_cols)] * max_cols
                data = [[_para_cell(c, body) for c in row] for row in pad_rows]
                t = Table(data, colWidths=col_widths, repeatRows=1)
                tbl_cmds: list[tuple[Any, ...]] = [
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#c8c8c8")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
                if pad_rows:
                    tbl_cmds.append(
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#ececec"))
                    )
                t.setStyle(TableStyle(tbl_cmds))
                story.append(t)
                story.append(Spacer(1, 0.2 * cm))
            continue

        mu = _RE_UL.match(s)
        if mu:
            flush_plain_pdf()
            inner = _strip_gfm_task_list_prefix(_strip_inline_md(mu.group(1)))
            txt = xml_escape(inner)
            story.append(Paragraph(txt, bullet_body, bulletText="\u2022"))
            i += 1
            continue

        mo = _RE_OL.match(s)
        if mo:
            flush_plain_pdf()
            n, rest = mo.group(1), mo.group(2)
            txt = xml_escape(_strip_inline_md(rest))
            story.append(Paragraph(f"{n}. {txt}", ol_body))
            i += 1
            continue

        mq = _RE_BLOCKQUOTE.match(s)
        if mq:
            flush_plain_pdf()
            inner = mq.group(1).strip()
            if inner:
                story.append(
                    Paragraph(xml_escape(_strip_inline_md(inner)), quote_style)
                )
            i += 1
            continue

        if _is_plain_markdown_line(s):
            plain_buf.append(s)
            i += 1
            continue

        flush_plain_pdf()
        story.append(Paragraph(xml_escape(_strip_inline_md(s)), body))
        i += 1

    flush_plain_pdf()
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )
    doc.build(story)
    return buf.getvalue()
