"""Markdown → Word（.docx）/ 简易 PDF；供任务报告与策略稿导出。"""
from __future__ import annotations

import os
import re
from io import BytesIO
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape as xml_escape


def _strip_inline_md(s: str) -> str:
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"`([^`]+)`", r"\1", s)
    return s


def _is_table_sep(line: str) -> bool:
    t = line.strip()
    if not t.startswith("|"):
        return False
    inner = t.strip("|").replace(" ", "")
    return bool(inner) and all(p in ("", "---", ":---", "---:", ":---:") for p in t.split("|"))


_img_line = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")


def markdown_to_docx_bytes(md: str, *, asset_root: Path | None = None) -> bytes:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Inches, Pt

    doc = Document()
    try:
        style = doc.styles["Normal"]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(10.5)
    except Exception:
        pass

    lines = (md or "").replace("\r\n", "\n").split("\n")
    i = 0
    in_fence = False
    while i < len(lines):
        raw = lines[i]
        if raw.strip().startswith("```"):
            in_fence = not in_fence
            i += 1
            continue
        if in_fence:
            p = doc.add_paragraph(xml_escape(raw) or " ")
            p.style = doc.styles["Normal"]
            for run in p.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            i += 1
            continue

        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            i += 1
            continue
        if line.startswith("# "):
            doc.add_heading(_strip_inline_md(line[2:].strip()), level=0)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(_strip_inline_md(line[3:].strip()), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(_strip_inline_md(line[4:].strip()), level=2)
            i += 1
            continue
        if line.startswith("#### "):
            doc.add_heading(_strip_inline_md(line[5:].strip()), level=3)
            i += 1
            continue
        mimg = _img_line.match(line.strip())
        if mimg and asset_root is not None:
            rel = mimg.group(2).strip()
            if not (rel.startswith("http://") or rel.startswith("https://")):
                img_path = (asset_root / rel).resolve()
                try:
                    img_path.relative_to(asset_root.resolve())
                except ValueError:
                    i += 1
                    continue
                if img_path.is_file():
                    doc.add_picture(str(img_path), width=Inches(5.9))
            i += 1
            continue
        if line.strip().startswith("|"):
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_line = lines[i].strip()
                if _is_table_sep(row_line):
                    i += 1
                    continue
                cells = [c.strip() for c in row_line.strip("|").split("|")]
                rows.append([_strip_inline_md(c) for c in cells])
                i += 1
            if rows:
                max_cols = max(len(r) for r in rows)
                pad_rows = [r + [""] * (max_cols - len(r)) for r in rows]
                tbl = doc.add_table(rows=len(pad_rows), cols=max_cols)
                tbl.style = "Table Grid"
                for ri, row in enumerate(pad_rows):
                    for ci, cell in enumerate(row):
                        tbl.rows[ri].cells[ci].text = cell
            continue

        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        text = _strip_inline_md(line)
        p.add_run(text)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _pdf_font_candidates() -> list[Path]:
    raw = (os.environ.get("MA_PDF_FONT") or "").strip()
    out: list[Path] = []
    if raw:
        out.append(Path(raw))
    windir = os.environ.get("WINDIR", r"C:\Windows")
    out.extend(
        [
            Path(windir) / "Fonts" / "simhei.ttf",
            Path(windir) / "Fonts" / "simsun.ttc",
            Path(windir) / "Fonts" / "msyh.ttf",
        ]
    )
    # Linux / 容器常见中文字体（路径不存在则跳过）
    out.extend(
        [
            Path("/usr/share/fonts/truetype/wqy/wqy-microhei.ttc"),
            Path("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"),
            Path("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc"),
            Path("/usr/share/fonts/truetype/noto/NotoSansCJKsc-Regular.otf"),
            Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
        ]
    )
    return out


def _pdf_flowable_image(img_path: Path, *, max_w: float, max_h: float) -> Any:
    """将插图缩放到不超过 max_w×max_h（ReportLab 单位，与 cm 一致），保持宽高比，避免矩阵长图撑爆版面。"""
    from reportlab.lib.utils import ImageReader
    from reportlab.platypus import Image as RLImage

    p = str(img_path)
    try:
        ir = ImageReader(p)
        iw, ih = ir.getSize()
    except Exception:
        return RLImage(p, width=max_w * 0.9, height=max_h * 0.9)
    if iw <= 0 or ih <= 0:
        return RLImage(p, width=max_w * 0.9, height=max_h * 0.9)
    w = float(max_w)
    h = w * (float(ih) / float(iw))
    if h > float(max_h):
        h = float(max_h)
        w = h * (float(iw) / float(ih))
    return RLImage(p, width=w, height=h)


def markdown_to_pdf_bytes(md: str, *, asset_root: Path | None = None) -> bytes:
    """简易纯文本流式 PDF；需本机 .ttf 中文字体或环境变量 MA_PDF_FONT。"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    font_name = "MaExportCJK"
    registered = False
    for p in _pdf_font_candidates():
        if not p.is_file():
            continue
        try:
            if p.suffix.lower() == ".ttc":
                try:
                    pdfmetrics.registerFont(
                        TTFont(font_name, str(p), subfontIndex=0)
                    )
                except TypeError:
                    pdfmetrics.registerFont(TTFont(font_name, str(p)))
            else:
                pdfmetrics.registerFont(TTFont(font_name, str(p)))
            registered = True
            break
        except Exception:
            continue
    if not registered:
        raise ValueError(
            "未找到可用的中文字体文件。请在 Windows 上安装黑体/宋体，"
            "或设置环境变量 MA_PDF_FONT 指向 .ttf 文件路径。"
        )

    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        name="BodyCJK",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=10,
        leading=14,
    )
    h1s = ParagraphStyle(
        name="H1CJK",
        parent=body,
        fontSize=16,
        leading=20,
        spaceAfter=8,
    )
    h2s = ParagraphStyle(
        name="H2CJK",
        parent=body,
        fontSize=13,
        leading=17,
        spaceAfter=6,
    )

    story: list[Any] = []
    lines = (md or "").replace("\r\n", "\n").split("\n")
    in_fence = False
    for raw in lines:
        if raw.strip().startswith("```"):
            in_fence = not in_fence
            continue
        s = raw.rstrip()
        if in_fence:
            story.append(Paragraph(xml_escape(s or " "), body))
            story.append(Spacer(1, 0.1 * cm))
            continue
        if not s.strip():
            story.append(Spacer(1, 0.15 * cm))
            continue
        mimg = _img_line.match(s.strip())
        if mimg and asset_root is not None:
            rel = mimg.group(2).strip()
            if not (rel.startswith("http://") or rel.startswith("https://")):
                img_path = (asset_root / rel).resolve()
                try:
                    img_path.relative_to(asset_root.resolve())
                except ValueError:
                    continue
                if img_path.is_file():
                    # 版面可用高度需小于正文框（A4 减边距后约 24.6cm），否则 ReportLab 报 LayoutError
                    story.append(
                        _pdf_flowable_image(
                            img_path, max_w=13 * cm, max_h=24 * cm
                        )
                    )
                    story.append(Spacer(1, 0.2 * cm))
            continue
        plain = _strip_inline_md(s)
        text = xml_escape(plain)
        if s.startswith("# "):
            story.append(Paragraph(xml_escape(plain[2:]), h1s))
        elif s.startswith("## "):
            story.append(Paragraph(xml_escape(plain[3:]), h2s))
        elif s.startswith("### "):
            story.append(Paragraph(xml_escape(plain[4:]), body))
        elif s.strip().startswith("|"):
            story.append(Paragraph(text.replace("|", " │ "), body))
        else:
            story.append(Paragraph(text, body))

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )
    doc.build(story)
    return buf.getvalue()
